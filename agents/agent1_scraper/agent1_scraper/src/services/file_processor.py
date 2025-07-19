"""
File Processor Service for Agent 1: Web Scraper & Data Collector
"""

import os
import io
import json
import threading
import mimetypes
from datetime import datetime
from typing import Dict, List, Any, Optional

try:
    import pandas as pd
    import numpy as np
except ImportError:
    pd = None
    np = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import cv2
except ImportError:
    cv2 = None

try:
    import magic
except ImportError:
    magic = None
class FileProcessorService:
    """Service for processing uploaded files"""
    
    def __init__(self):
        self.processing_threads = {}
        
    def start_processing(self, file_id: str):
        """Start processing a file asynchronously"""
        def process_file():
            from src.models.scraper import db, FileUpload
            
            try:
                # Get file from database
                file_upload = FileUpload.query.filter_by(file_id=file_id).first()
                if not file_upload:
                    return
                
                # Update status
                file_upload.processing_status = 'processing'
                db.session.commit()
                
                # Process file based on type
                result = self.process_file_by_type(file_upload.file_path, file_upload.mime_type)
                
                if result['success']:
                    file_upload.processing_status = 'completed'
                    file_upload.extracted_content = result['content']
                    file_upload.meta_data = result['metadata']
                else:
                    file_upload.processing_status = 'failed'
                    file_upload.error_message = result.get('error')
                
                file_upload.processed_at = datetime.utcnow()
                db.session.commit()
                
            except Exception as e:
                # Update with error
                file_upload = FileUpload.query.filter_by(file_id=file_id).first()
                if file_upload:
                    file_upload.processing_status = 'failed'
                    file_upload.error_message = str(e)
                    file_upload.processed_at = datetime.utcnow()
                    db.session.commit()
            
            finally:
                # Clean up thread reference
                if file_id in self.processing_threads:
                    del self.processing_threads[file_id]
        
        # Start processing in separate thread
        thread = threading.Thread(target=process_file, daemon=True)
        thread.start()
        self.processing_threads[file_id] = thread
    
    def process_file_by_type(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process file based on its MIME type"""
        try:
            if not os.path.exists(file_path):
                return {'success': False, 'error': 'File not found'}
            
            # Detect MIME type if not provided
            if not mime_type:
                mime_type = magic.from_file(file_path, mime=True)
            
            # Route to appropriate processor
            if mime_type.startswith('text/'):
                return self.process_text_file(file_path, mime_type)
            elif mime_type == 'application/pdf':
                return self.process_pdf_file(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'application/msword']:
                return self.process_word_file(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                              'application/vnd.ms-excel', 'text/csv']:
                return self.process_spreadsheet_file(file_path, mime_type)
            elif mime_type == 'application/json':
                return self.process_json_file(file_path)
            elif mime_type in ['application/xml', 'text/xml']:
                return self.process_xml_file(file_path)
            elif mime_type.startswith('image/'):
                return self.process_image_file(file_path, mime_type)
            elif mime_type.startswith('video/'):
                return self.process_video_file(file_path, mime_type)
            elif mime_type.startswith('audio/'):
                return self.process_audio_file(file_path, mime_type)
            else:
                return self.process_binary_file(file_path, mime_type)
                
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_text_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Basic text analysis
            lines = content.split('\n')
            words = content.split()
            
            metadata = {
                'file_type': 'text',
                'mime_type': mime_type,
                'encoding': 'utf-8',
                'line_count': len(lines),
                'word_count': len(words),
                'character_count': len(content),
                'processed_at': datetime.utcnow().isoformat()
            }
            
            # Extract keywords (simple implementation)
            word_freq = {}
            for word in words:
                word = word.lower().strip('.,!?;:"()[]{}')
                if len(word) > 2:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]
            metadata['top_words'] = top_words
            
            return {
                'success': True,
                'content': {
                    'text': content,
                    'lines': lines[:100],  # First 100 lines
                    'summary': content[:1000] + '...' if len(content) > 1000 else content
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Extract text from all pages
                text_content = ""
                page_texts = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    page_texts.append({
                        'page_number': page_num + 1,
                        'text': page_text
                    })
                    text_content += page_text + "\n"
                
                # Get PDF metadata
                pdf_info = pdf_reader.metadata
                
                metadata = {
                    'file_type': 'pdf',
                    'mime_type': 'application/pdf',
                    'page_count': len(pdf_reader.pages),
                    'character_count': len(text_content),
                    'word_count': len(text_content.split()),
                    'processed_at': datetime.utcnow().isoformat()
                }
                
                # Add PDF metadata if available
                if pdf_info:
                    metadata['pdf_info'] = {
                        'title': pdf_info.get('/Title'),
                        'author': pdf_info.get('/Author'),
                        'subject': pdf_info.get('/Subject'),
                        'creator': pdf_info.get('/Creator'),
                        'producer': pdf_info.get('/Producer'),
                        'creation_date': str(pdf_info.get('/CreationDate')),
                        'modification_date': str(pdf_info.get('/ModDate'))
                    }
                
                return {
                    'success': True,
                    'content': {
                        'text': text_content,
                        'pages': page_texts,
                        'summary': text_content[:1000] + '...' if len(text_content) > 1000 else text_content
                    },
                    'metadata': metadata
                }
                
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_word_file(self, file_path: str) -> Dict[str, Any]:
        """Process Word documents"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            metadata = {
                'file_type': 'word',
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'character_count': len(full_text),
                'word_count': len(full_text.split()),
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return {
                'success': True,
                'content': {
                    'text': full_text,
                    'paragraphs': paragraphs,
                    'tables': tables,
                    'summary': full_text[:1000] + '...' if len(full_text) > 1000 else full_text
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_spreadsheet_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process spreadsheet files (Excel, CSV)"""
        try:
            if mime_type == 'text/csv':
                df = pd.read_csv(file_path)
                sheets = {'Sheet1': df}
            else:
                # Excel file
                excel_file = pd.ExcelFile(file_path)
                sheets = {}
                for sheet_name in excel_file.sheet_names:
                    sheets[sheet_name] = pd.read_excel(excel_file, sheet_name=sheet_name)
            
            # Process each sheet
            processed_sheets = {}
            total_rows = 0
            total_cols = 0
            
            for sheet_name, df in sheets.items():
                # Convert to JSON-serializable format
                sheet_data = {
                    'columns': df.columns.tolist(),
                    'data': df.head(100).to_dict('records'),  # First 100 rows
                    'shape': df.shape,
                    'dtypes': df.dtypes.astype(str).to_dict(),
                    'summary': df.describe().to_dict() if df.select_dtypes(include='number').shape[1] > 0 else {}
                }
                
                processed_sheets[sheet_name] = sheet_data
                total_rows += df.shape[0]
                total_cols = max(total_cols, df.shape[1])
            
            metadata = {
                'file_type': 'spreadsheet',
                'mime_type': mime_type,
                'sheet_count': len(sheets),
                'total_rows': total_rows,
                'total_columns': total_cols,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return {
                'success': True,
                'content': {
                    'sheets': processed_sheets,
                    'summary': f"Spreadsheet with {len(sheets)} sheet(s), {total_rows} total rows"
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_json_file(self, file_path: str) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Analyze JSON structure
            def analyze_structure(obj, depth=0):
                if depth > 5:  # Limit recursion depth
                    return "..."
                
                if isinstance(obj, dict):
                    return {k: analyze_structure(v, depth+1) for k, v in list(obj.items())[:10]}
                elif isinstance(obj, list):
                    if len(obj) > 0:
                        return [analyze_structure(obj[0], depth+1)] + (["..."] if len(obj) > 1 else [])
                    return []
                else:
                    return type(obj).__name__
            
            structure = analyze_structure(data)
            
            metadata = {
                'file_type': 'json',
                'mime_type': 'application/json',
                'structure': structure,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            # Add size information
            if isinstance(data, dict):
                metadata['key_count'] = len(data)
            elif isinstance(data, list):
                metadata['item_count'] = len(data)
            
            return {
                'success': True,
                'content': {
                    'data': data,
                    'structure': structure,
                    'summary': f"JSON file with {type(data).__name__} root type"
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_xml_file(self, file_path: str) -> Dict[str, Any]:
        """Process XML files"""
        try:
            import xml.etree.ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Convert XML to dict (simplified)
            def xml_to_dict(element):
                result = {}
                
                # Add attributes
                if element.attrib:
                    result['@attributes'] = element.attrib
                
                # Add text content
                if element.text and element.text.strip():
                    if len(element) == 0:
                        return element.text.strip()
                    result['#text'] = element.text.strip()
                
                # Add child elements
                for child in element:
                    child_data = xml_to_dict(child)
                    if child.tag in result:
                        if not isinstance(result[child.tag], list):
                            result[child.tag] = [result[child.tag]]
                        result[child.tag].append(child_data)
                    else:
                        result[child.tag] = child_data
                
                return result
            
            data = {root.tag: xml_to_dict(root)}
            
            metadata = {
                'file_type': 'xml',
                'mime_type': 'application/xml',
                'root_tag': root.tag,
                'namespace': root.tag.split('}')[0][1:] if '}' in root.tag else None,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return {
                'success': True,
                'content': {
                    'data': data,
                    'summary': f"XML file with root element '{root.tag}'"
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_image_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process image files"""
        try:
            # Open image with PIL
            with Image.open(file_path) as img:
                # Basic image info
                width, height = img.size
                mode = img.mode
                format_name = img.format
                
                # Get EXIF data if available
                exif_data = {}
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    for tag_id, value in exif.items():
                        tag = Image.ExifTags.TAGS.get(tag_id, tag_id)
                        exif_data[tag] = str(value)
                
                metadata = {
                    'file_type': 'image',
                    'mime_type': mime_type,
                    'width': width,
                    'height': height,
                    'mode': mode,
                    'format': format_name,
                    'exif': exif_data,
                    'processed_at': datetime.utcnow().isoformat()
                }
                
                return {
                    'success': True,
                    'content': {
                        'summary': f"Image: {width}x{height} pixels, {mode} mode",
                        'dimensions': {'width': width, 'height': height},
                        'exif': exif_data
                    },
                    'metadata': metadata
                }
                
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_video_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process video files"""
        try:
            # Use OpenCV to get video info
            cap = cv2.VideoCapture(file_path)
            
            if not cap.isOpened():
                return {'success': False, 'error': 'Could not open video file'}
            
            # Get video properties
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            fps = cap.get(cv2.CAP_PROP_FPS)
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0
            
            cap.release()
            
            metadata = {
                'file_type': 'video',
                'mime_type': mime_type,
                'width': width,
                'height': height,
                'frame_count': frame_count,
                'fps': fps,
                'duration_seconds': duration,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return {
                'success': True,
                'content': {
                    'summary': f"Video: {width}x{height}, {duration:.1f}s duration, {fps:.1f} FPS",
                    'dimensions': {'width': width, 'height': height},
                    'duration': duration,
                    'frame_rate': fps
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_audio_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process audio files"""
        try:
            # Basic file info
            file_size = os.path.getsize(file_path)
            
            metadata = {
                'file_type': 'audio',
                'mime_type': mime_type,
                'file_size': file_size,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return {
                'success': True,
                'content': {
                    'summary': f"Audio file: {mime_type}, {file_size} bytes",
                    'file_size': file_size
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_binary_file(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process binary files (fallback)"""
        try:
            file_size = os.path.getsize(file_path)
            
            metadata = {
                'file_type': 'binary',
                'mime_type': mime_type,
                'file_size': file_size,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return {
                'success': True,
                'content': {
                    'summary': f"Binary file: {mime_type}, {file_size} bytes",
                    'file_size': file_size
                },
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}

