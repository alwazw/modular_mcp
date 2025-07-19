"""
Document Processor Service for Agent 2: Knowledge Base Creator
"""

import uuid
import threading
import os
import tempfile
import hashlib
from datetime import datetime
from typing import Dict, List, Any, Optional
import requests
from src.models.knowledge import db, Document, KnowledgeBase
from src.services.knowledge_processor import KnowledgeProcessorService

class DocumentProcessorService:
    """Service for processing documents from various sources"""
    
    def __init__(self):
        self.processing_threads = {}
        self.knowledge_processor = KnowledgeProcessorService()
        
        # Agent 1 configuration
        self.agent1_base_url = os.getenv('AGENT1_BASE_URL', 'http://localhost:5000')
    
    def start_document_processing(self, document_id: str, kb_config: Dict[str, Any]):
        """Start document processing asynchronously"""
        def process_document():
            try:
                # Get document from database
                document = Document.query.filter_by(document_id=document_id).first()
                if not document:
                    return
                
                # Update status
                document.processing_status = 'processing'
                db.session.commit()
                
                # Process based on source type
                if document.source_type == 'file':
                    result = self.process_file_document(document, kb_config)
                elif document.source_type == 'url':
                    result = self.process_url_document(document, kb_config)
                elif document.source_type == 'text':
                    result = self.process_text_document(document, kb_config)
                elif document.source_type == 'agent1':
                    result = self.process_agent1_document(document, kb_config)
                else:
                    result = {'success': False, 'error': f'Unknown source type: {document.source_type}'}
                
                # Update document with results
                if result['success']:
                    document.processing_status = 'completed'
                    document.processed_at = datetime.utcnow()
                    if 'chunk_count' in result:
                        document.chunk_count = result['chunk_count']
                else:
                    document.processing_status = 'failed'
                    document.error_message = result.get('error')
                
                db.session.commit()
                
            except Exception as e:
                # Update with error
                document = Document.query.filter_by(document_id=document_id).first()
                if document:
                    document.processing_status = 'failed'
                    document.error_message = str(e)
                    db.session.commit()
            
            finally:
                # Clean up thread reference
                if document_id in self.processing_threads:
                    del self.processing_threads[document_id]
        
        # Start processing in separate thread
        thread = threading.Thread(target=process_document, daemon=True)
        thread.start()
        self.processing_threads[document_id] = thread
    
    def start_file_processing(self, document_id: str, file_path: str, kb_config: Dict[str, Any]):
        """Start file processing asynchronously"""
        def process_file():
            try:
                # Get document from database
                document = Document.query.filter_by(document_id=document_id).first()
                if not document:
                    return
                
                # Update status
                document.processing_status = 'processing'
                db.session.commit()
                
                # Read file content
                content = self.extract_file_content(file_path, document.content_type)
                
                if not content:
                    document.processing_status = 'failed'
                    document.error_message = 'Could not extract content from file'
                    db.session.commit()
                    return
                
                # Update document with content
                document.raw_content = content
                document.processed_content = self.clean_content(content)
                document.content_hash = hashlib.sha256(content.encode()).hexdigest()
                
                # Process into chunks
                result = self.knowledge_processor.process_document_by_id(document_id, kb_config)
                
                # Clean up temporary file
                try:
                    os.unlink(file_path)
                except:
                    pass
                
            except Exception as e:
                # Update with error
                document = Document.query.filter_by(document_id=document_id).first()
                if document:
                    document.processing_status = 'failed'
                    document.error_message = str(e)
                    db.session.commit()
            
            finally:
                # Clean up thread reference
                if document_id in self.processing_threads:
                    del self.processing_threads[document_id]
        
        # Start processing in separate thread
        thread = threading.Thread(target=process_file, daemon=True)
        thread.start()
        self.processing_threads[document_id] = thread
    
    def start_agent1_import(self, document_id: str, agent1_content: Dict[str, Any], kb_config: Dict[str, Any]):
        """Start Agent 1 content import asynchronously"""
        def import_content():
            try:
                # Get document from database
                document = Document.query.filter_by(document_id=document_id).first()
                if not document:
                    return
                
                # Update status
                document.processing_status = 'processing'
                db.session.commit()
                
                # Process the imported content
                result = self.knowledge_processor.process_document_by_id(document_id, kb_config)
                
            except Exception as e:
                # Update with error
                document = Document.query.filter_by(document_id=document_id).first()
                if document:
                    document.processing_status = 'failed'
                    document.error_message = str(e)
                    db.session.commit()
            
            finally:
                # Clean up thread reference
                if document_id in self.processing_threads:
                    del self.processing_threads[document_id]
        
        # Start processing in separate thread
        thread = threading.Thread(target=import_content, daemon=True)
        thread.start()
        self.processing_threads[document_id] = thread
    
    def process_file_document(self, document: Document, kb_config: Dict[str, Any]) -> Dict[str, Any]:
        """Process a file document"""
        try:
            if not document.source_path or not os.path.exists(document.source_path):
                return {'success': False, 'error': 'File not found'}
            
            # Extract content from file
            content = self.extract_file_content(document.source_path, document.content_type)
            
            if not content:
                return {'success': False, 'error': 'Could not extract content from file'}
            
            # Update document with content
            document.raw_content = content
            document.processed_content = self.clean_content(content)
            document.content_hash = hashlib.sha256(content.encode()).hexdigest()
            db.session.commit()
            
            # Process into chunks
            return self.knowledge_processor.process_document_by_id(document.document_id, kb_config)
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_url_document(self, document: Document, kb_config: Dict[str, Any]) -> Dict[str, Any]:
        """Process a URL document by scraping with Agent 1"""
        try:
            if not document.source_path:
                return {'success': False, 'error': 'No URL provided'}
            
            # Request Agent 1 to scrape the URL
            scrape_result = self.request_agent1_scraping(document.source_path)
            
            if not scrape_result['success']:
                return {'success': False, 'error': f'Agent 1 scraping failed: {scrape_result.get("error")}'}
            
            # Extract content from scraping result
            scraped_data = scrape_result['data']
            content = scraped_data.get('cleaned_content') or scraped_data.get('raw_content', '')
            
            if not content:
                return {'success': False, 'error': 'No content extracted from URL'}
            
            # Update document with content
            document.raw_content = scraped_data.get('raw_content')
            document.processed_content = content
            document.content_hash = hashlib.sha256(content.encode()).hexdigest()
            document.meta_data = {
                **document.meta_data,
                'scraped_title': scraped_data.get('title'),
                'scraped_at': datetime.utcnow().isoformat(),
                'agent1_metadata': scraped_data.get('metadata', {})
            }
            db.session.commit()
            
            # Process into chunks
            return self.knowledge_processor.process_document_by_id(document.document_id, kb_config)
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_text_document(self, document: Document, kb_config: Dict[str, Any]) -> Dict[str, Any]:
        """Process a text document"""
        try:
            content = document.raw_content
            
            if not content:
                return {'success': False, 'error': 'No content provided'}
            
            # Clean and process content
            document.processed_content = self.clean_content(content)
            document.content_hash = hashlib.sha256(content.encode()).hexdigest()
            db.session.commit()
            
            # Process into chunks
            return self.knowledge_processor.process_document_by_id(document.document_id, kb_config)
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def process_agent1_document(self, document: Document, kb_config: Dict[str, Any]) -> Dict[str, Any]:
        """Process a document imported from Agent 1"""
        try:
            # Content should already be set from Agent 1
            content = document.processed_content or document.raw_content
            
            if not content:
                return {'success': False, 'error': 'No content from Agent 1'}
            
            # Process into chunks
            return self.knowledge_processor.process_document_by_id(document.document_id, kb_config)
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def extract_file_content(self, file_path: str, content_type: str) -> Optional[str]:
        """Extract text content from various file types"""
        try:
            if not os.path.exists(file_path):
                return None
            
            # Handle different file types
            if content_type and content_type.startswith('text/'):
                # Plain text files
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif content_type == 'application/pdf':
                # PDF files
                return self.extract_pdf_content(file_path)
            
            elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                'application/msword']:
                # Word documents
                return self.extract_word_content(file_path)
            
            elif content_type == 'application/json':
                # JSON files
                with open(file_path, 'r', encoding='utf-8') as f:
                    import json
                    data = json.load(f)
                    return json.dumps(data, indent=2)
            
            else:
                # Try to read as text
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
                except:
                    return None
            
        except Exception as e:
            print(f"Error extracting file content: {e}")
            return None
    
    def extract_pdf_content(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_content = ""
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                return text_content
                
        except Exception as e:
            print(f"Error extracting PDF content: {e}")
            return None
    
    def extract_word_content(self, file_path: str) -> Optional[str]:
        """Extract text from Word document"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return text_content
            
        except Exception as e:
            print(f"Error extracting Word content: {e}")
            return None
    
    def clean_content(self, content: str) -> str:
        """Clean and normalize content"""
        if not content:
            return ""
        
        import re
        
        # Remove excessive whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove HTML tags if present
        content = re.sub(r'<[^>]+>', '', content)
        
        # Normalize line breaks
        content = re.sub(r'\r\n|\r|\n', '\n', content)
        
        # Remove excessive line breaks
        content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
        
        return content.strip()
    
    def request_agent1_scraping(self, url: str) -> Dict[str, Any]:
        """Request Agent 1 to scrape a URL"""
        try:
            # Make request to Agent 1 scraping endpoint
            response = requests.post(
                f"{self.agent1_base_url}/api/scraper/scrape",
                json={
                    'url': url,
                    'wait_for_content': True,
                    'extract_text': True
                },
                timeout=60
            )
            
            if response.status_code == 200:
                return {
                    'success': True,
                    'data': response.json()
                }
            else:
                return {
                    'success': False,
                    'error': f'Agent 1 returned status {response.status_code}: {response.text}'
                }
                
        except requests.exceptions.RequestException as e:
            return {
                'success': False,
                'error': f'Failed to connect to Agent 1: {str(e)}'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error requesting Agent 1 scraping: {str(e)}'
            }
    
    def fetch_agent1_content(self, content_id: str) -> Dict[str, Any]:
        """Fetch processed content from Agent 1"""
        try:
            # Make request to Agent 1 to get processed content
            response = requests.get(
                f"{self.agent1_base_url}/api/scraper/content/{content_id}",
                timeout=30
            )
            
            if response.status_code == 200:
                return {
                    'success': True,
                    'data': response.json()
                }
            else:
                return {
                    'success': False,
                    'error': f'Agent 1 returned status {response.status_code}: {response.text}'
                }
                
        except requests.exceptions.RequestException as e:
            return {
                'success': False,
                'error': f'Failed to connect to Agent 1: {str(e)}'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error fetching Agent 1 content: {str(e)}'
            }
    
    def sync_with_agent1(self, kb_id: str) -> Dict[str, Any]:
        """Sync knowledge base with Agent 1 scraped content"""
        try:
            # Get knowledge base
            kb = KnowledgeBase.query.filter_by(kb_id=kb_id).first()
            if not kb:
                return {'success': False, 'error': 'Knowledge base not found'}
            
            # Request list of scraped content from Agent 1
            response = requests.get(
                f"{self.agent1_base_url}/api/scraper/content",
                params={'status': 'completed'},
                timeout=30
            )
            
            if response.status_code != 200:
                return {
                    'success': False,
                    'error': f'Failed to get content list from Agent 1: {response.status_code}'
                }
            
            content_list = response.json().get('content', [])
            imported_count = 0
            
            # Import each piece of content
            for content_item in content_list:
                # Check if already imported
                existing_doc = Document.query.filter_by(
                    content_hash=content_item.get('content_hash')
                ).first()
                
                if existing_doc:
                    continue  # Skip if already imported
                
                # Create document from Agent 1 content
                document = Document(
                    document_id=str(uuid.uuid4()),
                    title=content_item.get('title', 'Imported from Agent 1'),
                    source_type='agent1',
                    source_path=content_item.get('url'),
                    content_type='text/html',
                    raw_content=content_item.get('raw_content'),
                    processed_content=content_item.get('cleaned_content'),
                    meta_data={
                        'agent1_content_id': content_item.get('id'),
                        'imported_at': datetime.utcnow().isoformat(),
                        'agent1_metadata': content_item.get('metadata', {})
                    },
                    content_hash=content_item.get('content_hash')
                )
                
                db.session.add(document)
                db.session.commit()
                
                # Start processing
                self.start_document_processing(document.document_id, kb.to_dict())
                imported_count += 1
            
            return {
                'success': True,
                'imported_count': imported_count,
                'total_available': len(content_list)
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}

